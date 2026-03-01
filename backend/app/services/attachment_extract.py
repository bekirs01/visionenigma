"""
Извлечение текста из вложений (PDF, изображения) для передачи в AI.
- PDF: pypdf text extraction, при пустом — OCR (pdf2image + pytesseract).
- Изображения (jpg/png/webp): pytesseract OCR.
- Видео: не анализируем, возвращаем заметку для AI.
"""
from typing import Tuple

# Максимум символов извлечённого текста для промпта (без обрезки с "...")
EXTRACTED_TEXT_MAX_CHARS = 30_000
# PDF OCR: максимум страниц (большие документы не обрабатываем полностью)
PDF_OCR_MAX_PAGES = 10


def extract_text_from_attachment(filename: str, mime_type: str, data: bytes) -> Tuple[bool, str]:
    """
    Извлекает текст из вложения.

    Returns:
        (success, text_or_note)
        - success=True: text — извлечённый текст (может быть пустым если не удалось прочитать).
        - success=False: text — краткая заметка для AI (например "видео получено, оператор проверит").
    """
    mime = (mime_type or "").lower()
    fn = (filename or "").lower()

    # PDF
    if "pdf" in mime or fn.endswith(".pdf"):
        return _extract_pdf(data)

    # Изображения
    if any(x in mime for x in ("image/jpeg", "image/jpg", "image/png", "image/webp")):
        return _extract_image(data)

    # DOCX
    if fn.endswith(".docx") or "wordprocessingml" in mime:
        return _extract_docx(data)

    # DOC (eski format) - cikaramiyoruz ama bilgi veriyoruz
    if fn.endswith(".doc") or "msword" in mime:
        return False, "Получен файл в формате .doc. Для автоматического анализа рекомендуем формат .docx или PDF."

    # Видео — не анализируем
    if "video/" in mime or fn.endswith((".mp4", ".webm", ".mov", ".avi")):
        return False, "Получено видео-вложение. Анализ видео не выполняется; оператор просмотрит вручную."

    # Остальные типы (xls и т.д.) — не извлекаем
    return False, "Вложение получено, но автоматическое извлечение текста для данного формата не поддерживается. Попросите клиента описать содержание текстом или прислать PDF/изображение."


def _extract_docx(data: bytes) -> Tuple[bool, str]:
    """DOCX: python-docx ile paragraf metinlerini cikarir."""
    try:
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            t = para.text
            if t and t.strip():
                parts.append(t.strip())
        # Tablolardan da metin cikar
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts) if parts else ""
        if text.strip():
            return True, _truncate_safe(text.strip())
        return True, ""
    except Exception:
        return False, "DOCX dosyasi acilamadi. Lutfen icerigi metin olarak gonderin."


def _extract_pdf(data: bytes) -> Tuple[bool, str]:
    """PDF: сначала текст, при пустом — OCR (макс PDF_OCR_MAX_PAGES страниц)."""
    text = _pdf_text_extract(data)
    if (text or "").strip():
        return True, _truncate_safe(text.strip())
    # OCR fallback
    try:
        ocr_text = _pdf_ocr(data)
        if (ocr_text or "").strip():
            return True, _truncate_safe(ocr_text.strip())
    except Exception:
        pass
    # OCR недоступен или не сработал — AI всё равно получит пояснение и сформирует ответ
    return True, "[PDF görüntü tabanlı; OCR yapılamadı. Lütfen içeriği metin olarak yazın veya farklı format gönderin.]"


def _pdf_text_extract(data: bytes) -> str:
    """Извлечение текста из PDF через pypdf (или pdfminer.six)."""
    try:
        from pypdf import PdfReader
        from io import BytesIO
        reader = PdfReader(BytesIO(data))
        parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n\n".join(parts) if parts else ""
    except Exception:
        try:
            from pdfminer.high_level import extract_text
            from io import BytesIO
            return extract_text(BytesIO(data)) or ""
        except Exception:
            return ""


def _pdf_ocr(data: bytes) -> str:
    """OCR по страницам PDF (pdf2image + pytesseract). Максимум PDF_OCR_MAX_PAGES страниц."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        images = convert_from_bytes(data, dpi=150)
        parts = []
        for i, img in enumerate(images):
            if i >= PDF_OCR_MAX_PAGES:
                break
            text = pytesseract.image_to_string(img, lang="rus+eng")
            if text and text.strip():
                parts.append(text.strip())
        return "\n\n".join(parts) if parts else ""
    except Exception:
        return ""


def _extract_image(data: bytes) -> Tuple[bool, str]:
    """Изображение: pytesseract OCR."""
    try:
        import pytesseract
        from PIL import Image
        from io import BytesIO
        img = Image.open(BytesIO(data))
        text = pytesseract.image_to_string(img, lang="rus+eng")
        if (text or "").strip():
            return True, _truncate_safe(text.strip())
        return True, ""
    except Exception:
        return True, ""


def _truncate_safe(s: str, max_chars: int = EXTRACTED_TEXT_MAX_CHARS) -> str:
    """Обрезает до max_chars символов по границе слова, без добавления '...'."""
    s = s.strip()
    if len(s) <= max_chars:
        return s
    cut = s[:max_chars]
    last_space = cut.rfind(" ")
    if last_space > max_chars // 2:
        return cut[:last_space].strip()
    return cut.strip()
